# txt_to_docx.py
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import argparse

parser = argparse.ArgumentParser()
parser.add_argument('--input_dir', default='txt/pdf_txt')
parser.add_argument('--output_dir', default='word')
args = parser.parse_args()

def setup_document(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(3)

    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)


def add_line(doc: Document, line: str):
    text = line.rstrip()

    if text == "":
        doc.add_paragraph("")
        return

    if text[0].isdigit() and ". " in text[:6]:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        return

    if len(text) >= 2 and text[1] == ")":
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Cm(0.63)
        return

    doc.add_paragraph(text)


def txt_to_word(input_dir=args.input_dir, output_dir=args.output_dir):
    input_dir = Path(input_dir)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    for txt_file in sorted(input_dir.glob("*.txt")):
        doc = Document()
        setup_document(doc)

        lines = txt_file.read_text(encoding="utf-8").splitlines()
        for line in lines:
            add_line(doc, line)

        out = output_dir / f"{txt_file.stem}.docx"
        doc.save(out)
        print(f"Hazırdır ✅ → {out}")


if __name__ == "__main__":
    txt_to_word()